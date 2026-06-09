#!/usr/bin/env python3
"""
Image-to-DOCX converter using PaddleOCR Online API (PaddleOCR-VL-1.6).
Preserves layout: headings, paragraphs, tables, bold/italic formatting.

Usage: python3 paddle_api_convert.py <input_image> <output.docx>
"""

import json
import os
import re
import sys
import time

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---- Config ----
JOB_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
TOKEN = "11c8da827944176a9aa5e4efa998c04985e927d3"
MODEL = "PaddleOCR-VL-1.6"
HEADERS = {"Authorization": f"bearer {TOKEN}"}
OPTIONAL_PAYLOAD = {
    "useDocOrientationClassify": False,
    "useDocUnwarping": False,
    "useChartRecognition": False,
}


def submit_job(file_path):
    """Submit an OCR job to PaddleOCR API, return job ID."""
    data = {"model": MODEL, "optionalPayload": json.dumps(OPTIONAL_PAYLOAD)}
    with open(file_path, "rb") as f:
        files = {"file": f}
        resp = requests.post(JOB_URL, headers=HEADERS, data=data, files=files)
    if resp.status_code != 200:
        raise RuntimeError(f"API submit failed: {resp.status_code} {resp.text}")
    return resp.json()["data"]["jobId"]


def poll_job(job_id, max_wait=300):
    """Poll job status until done, return JSONL URL."""
    for _ in range(max_wait // 5):
        time.sleep(5)
        resp = requests.get(f"{JOB_URL}/{job_id}", headers=HEADERS)
        data = resp.json()["data"]
        state = data["state"]
        if state == "done":
            return data["resultUrl"]["jsonUrl"]
        elif state == "failed":
            raise RuntimeError(f"Job failed: {data.get('errorMsg', 'unknown')}")
    raise TimeoutError("Job timed out")


def fetch_markdown(jsonl_url):
    """Download JSONL and return combined markdown text."""
    resp = requests.get(jsonl_url)
    resp.raise_for_status()
    lines = resp.text.strip().split("\n")
    md_parts = []
    for line in lines:
        if not line.strip():
            continue
        parsed = json.loads(line)
        for lr in parsed["result"]["layoutParsingResults"]:
            md_parts.append(lr["markdown"]["text"])
    return "\n\n".join(md_parts)


def md_to_docx(markdown_text, docx_path):
    """Convert Markdown (with HTML tables) to a .docx file."""
    import markdown

    # Convert markdown to HTML
    html = markdown.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "codehilite"],
    )
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for element in soup.children:
        _process_element(doc, element)

    doc.save(docx_path)


def _process_element(doc, element, list_level=0):
    """Recursively process HTML elements and add to DOCX."""
    if element.name is None:
        # Text node
        text = str(element).strip()
        if text:
            p = doc.add_paragraph(text)
            p.style = doc.styles["Normal"]
        return

    tag = element.name

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        text = element.get_text(strip=True)
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            sizes = {1: Pt(22), 2: Pt(18), 3: Pt(14), 4: Pt(12), 5: Pt(11), 6: Pt(10)}
            run.font.size = sizes.get(level, Pt(11))
            if level <= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    elif tag == "p":
        _add_paragraph_with_inline(doc, element)

    elif tag == "table":
        _add_table(doc, element)

    elif tag in ("ul", "ol"):
        for li in element.find_all("li", recursive=False):
            prefix = "• " if tag == "ul" else f"{list_level + 1}. "
            p = doc.add_paragraph()
            run = p.add_run(prefix)
            for child in li.children:
                if child.name is None:
                    run_text = p.add_run(str(child))
                elif child.name in ("strong", "b"):
                    run_text = p.add_run(child.get_text())
                    run_text.bold = True
                elif child.name in ("em", "i"):
                    run_text = p.add_run(child.get_text())
                    run_text.italic = True
                else:
                    _process_inline(p, child)

    elif tag == "pre":
        text = element.get_text()
        if text.strip():
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    elif tag == "hr":
        doc.add_paragraph("─" * 60)

    elif tag == "blockquote":
        for child in element.children:
            if child.name:
                _process_element(doc, child)

    elif tag == "div":
        for child in element.children:
            _process_element(doc, child)


def _add_paragraph_with_inline(doc, p_element):
    """Add a paragraph with inline formatting (bold, italic, etc.)."""
    text = p_element.get_text(strip=True)
    if not text:
        return

    p = doc.add_paragraph()
    for child in p_element.children:
        if child.name is None:
            p.add_run(str(child))
        elif child.name in ("strong", "b"):
            r = p.add_run(child.get_text())
            r.bold = True
        elif child.name in ("em", "i"):
            r = p.add_run(child.get_text())
            r.italic = True
        elif child.name == "code":
            r = p.add_run(child.get_text())
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        elif child.name == "br":
            p.add_run("\n")
        elif child.name == "a":
            r = p.add_run(child.get_text())
            r.font.color.rgb = RGBColor(0, 0, 255)
            r.underline = True
        elif child.name == "img":
            p.add_run("[Image]")
        else:
            p.add_run(child.get_text() if hasattr(child, "get_text") else str(child))


def _add_table(doc, table_element):
    """Convert HTML table to DOCX table."""
    rows = table_element.find_all("tr")
    if not rows:
        return

    # Determine column count from first row
    first_row_cells = rows[0].find_all(["td", "th"])
    col_count = len(first_row_cells)
    if col_count == 0:
        return

    table = doc.add_table(rows=len(rows), cols=col_count, style="Table Grid")

    for row_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        for col_idx, cell in enumerate(cells):
            if col_idx >= col_count:
                break
            doc_cell = table.cell(row_idx, col_idx)
            text = cell.get_text(strip=True)
            doc_cell.text = text

            # Style header row
            is_header = cell.name == "th"
            if is_header:
                for paragraph in doc_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Center-align cells with "text-align: center"
            style_attr = cell.get("style", "")
            if "text-align: center" in style_attr:
                for paragraph in doc_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after table
    doc.add_paragraph()


def _process_inline(paragraph, element):
    """Add inline formatted text to an existing paragraph."""
    if element.name is None:
        paragraph.add_run(str(element))
    elif element.name in ("strong", "b"):
        r = paragraph.add_run(element.get_text())
        r.bold = True
    elif element.name in ("em", "i"):
        r = paragraph.add_run(element.get_text())
        r.italic = True
    elif element.name == "code":
        r = paragraph.add_run(element.get_text())
        r.font.name = "Courier New"
    elif element.name == "br":
        paragraph.add_run("\n")
    else:
        paragraph.add_run(element.get_text() if hasattr(element, "get_text") else str(element))


def convert_image_to_docx(img_path, docx_path):
    """Main conversion function: image → PaddleOCR API → Markdown → DOCX."""
    print(f"Submitting job for: {img_path}")
    job_id = submit_job(img_path)
    print(f"Job ID: {job_id}")

    print("Waiting for OCR results...")
    jsonl_url = poll_job(job_id)
    print(f"Results ready, downloading...")

    markdown_text = fetch_markdown(jsonl_url)
    print(f"Markdown length: {len(markdown_text)} chars")

    print("Converting to DOCX...")
    md_to_docx(markdown_text, docx_path)
    print(f"Saved: {docx_path}")

    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 paddle_api_convert.py <input_image> <output.docx>", file=sys.stderr)
        sys.exit(1)

    img_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.exists(img_path):
        print(f"ERROR: input file not found: {img_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = convert_image_to_docx(img_path, docx_path)
        if os.path.exists(result):
            print(f"OK: {result}")
        else:
            print("ERROR: output file not created", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

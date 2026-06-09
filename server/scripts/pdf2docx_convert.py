#!/usr/bin/env python3
"""PDF to DOCX converter using pdf2docx with optimized settings."""
import sys
import os

def convert(pdf_path, docx_path):
    from pdf2docx import Converter
    
    cv = Converter(pdf_path)
    cv.convert(
        docx_path,
        multi_processing=False,
        min_paragraph_height=3.0,
    )
    cv.close()
    
    # ---- Post-processing: fix number formatting in tables ----
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(docx_path)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Fix cell vertical alignment
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    
                    # Detect numeric cells and right-align
                    text = p.text.strip()
                    if text and _is_numeric_cell(text):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(docx_path)
    return docx_path


def _is_numeric_cell(text):
    """Check if cell content is primarily numeric (with formatting)."""
    # Remove common formatting chars
    cleaned = text.replace(',', '').replace('%', '').replace('$', '') \
                   .replace('¥', '').replace(' ', '').replace('\u00b1', '') \
                   .replace('+', '').replace('-', '').replace('.', '')
    # Also try digit-only check
    try:
        float(cleaned)
        return True
    except ValueError:
        pass
    return False


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 pdf2docx_convert.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    try:
        result = convert(pdf_path, docx_path)
        if os.path.exists(result):
            print(f"OK: {result}")
        else:
            print("ERROR: output file not created", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

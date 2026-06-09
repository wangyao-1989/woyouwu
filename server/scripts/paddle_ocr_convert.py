#!/usr/bin/env python3
"""
Image-to-DOCX converter using PaddleOCR PP-StructureV3 for professional layout recovery.
Usage: python3 paddle_ocr_convert.py <input_image> <output.docx>
"""
import sys
import os
import json
import tempfile

# Suppress model download noise
os.environ['PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK'] = 'True'


def convert_image_to_docx(img_path, docx_path):
    from paddleocr import PPStructureV3
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print('Loading PP-StructureV3 pipeline...')
    pipeline = PPStructureV3(
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_formula_recognition=False,  # Skip formula model (not downloaded)
    )

    print(f'Processing image: {img_path}')
    output = pipeline.predict(img_path)

    # ---- Build DOCX from structured output ----
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.size = Pt(11)

    for item in output:
        item_type = item.get('type', 'text')
        text = item.get('text', '').strip()
        bbox = item.get('bbox', [])

        if not text:
            continue

        if item_type == 'table':
            # Add table
            _add_table(doc, item)
        elif item_type == 'figure' or item_type == 'image':
            # Skip images for now, just add placeholder
            doc.add_paragraph('[Image]').alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif item_type == 'formula':
            doc.add_paragraph(text, style='Normal')
        else:
            # Text/paragraph/title etc.
            _add_text_paragraph(doc, item)

    doc.save(docx_path)
    print(f'Saved: {docx_path}')
    return docx_path


def _add_text_paragraph(doc, item):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = item.get('text', '').strip()
    if not text:
        return

    bbox = item.get('bbox', [])
    img_width = item.get('_img_width', 800)

    p = doc.add_paragraph()
    run = p.add_run(text)

    # Detect heading: short text + high Y confidence → likely title
    lines = text.split('\n')
    is_title = len(text) < 80 and len(lines) == 1
    if is_title:
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(11)

    # Left alignment by default


def _add_table(doc, table_item):
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cells = table_item.get('cells', [])
    if not cells:
        return

    # Parse cell grid: "row,column" format
    cell_grid = {}
    max_row = 0
    max_col = 0

    # PaddleOCR cells format: [[text, "row,col", rowspan, colspan], ...]
    # Or: {"text": "...", "bbox": [...], "cell_location": "row,col"}
    for cell_data in cells:
        if isinstance(cell_data, list):
            text = cell_data[0] if len(cell_data) > 0 else ''
            loc = cell_data[1] if len(cell_data) > 1 else '0,0'
        elif isinstance(cell_data, dict):
            text = cell_data.get('text', '')
            loc = cell_data.get('cell_location', '0,0')
        else:
            continue

        parts = loc.split(',')
        row = int(parts[0]) if len(parts) > 0 else 0
        col = int(parts[1]) if len(parts) > 1 else 0
        max_row = max(max_row, row)
        max_col = max(max_col, col)
        cell_grid[(row, col)] = text

    if max_row == 0 and max_col == 0:
        # Fallback: just add text lines
        for cell_data in cells:
            text = cell_data[0] if isinstance(cell_data, list) else cell_data.get('text', '')
            if text.strip():
                doc.add_paragraph(text.strip())
        return

    # Create table
    rows = max_row + 1
    cols = max_col + 1
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')

    for (r, c), text in cell_grid.items():
        cell = table.cell(r, c)
        cell.text = text.strip()
        # Make header row bold
        if r == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph()  # spacing after table


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 paddle_ocr_convert.py <input_image> <output.docx>", file=sys.stderr)
        sys.exit(1)

    img_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.exists(img_path):
        print(f"ERROR: input file not found: {img_path}", file=sys.stderr)
        sys.exit(1)

    try:
        result = convert_image_to_docx(img_path, docx_path)
        if os.path.exists(result):
            print(f"OK: {result}")
        else:
            print("ERROR: output file not created", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
